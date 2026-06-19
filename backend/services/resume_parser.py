"""Local PDF / DOCX text extraction + regex section parsing.

Designed to be cheap, deterministic, and fully offline.  AI enrichment is a
separate step (services don't call any LLM here).
"""

from __future__ import annotations

import io
import logging
import re
from typing import Iterable, List, Optional, Tuple

logger = logging.getLogger(__name__)

# A pragmatic, expandable skills lexicon. Used for skill mining only; new tokens
# captured via the heuristic SECTION scanner are also returned.
SKILLS_LEXICON = {
    # programming
    "python", "java", "javascript", "typescript", "c++", "c#", "go", "rust", "ruby", "php", "swift",
    "kotlin", "scala", "r", "matlab", "bash", "shell", "sql", "html", "css", "react", "angular",
    "vue", "next.js", "node", "node.js", "express", "django", "flask", "fastapi", "spring",
    # data / ml
    "pandas", "numpy", "scikit-learn", "sklearn", "tensorflow", "pytorch", "keras", "spark",
    "hadoop", "airflow", "snowflake", "bigquery", "redshift", "tableau", "power bi", "looker",
    "matplotlib", "seaborn", "etl", "data warehousing", "statistics",
    # qa / automation
    "selenium", "playwright", "cypress", "pytest", "junit", "testng", "appium",
    "manual testing", "automation testing", "regression testing", "api testing", "jira", "postman",
    # cloud / devops
    "aws", "azure", "gcp", "docker", "kubernetes", "terraform", "ansible", "jenkins", "github actions",
    "ci/cd", "linux",
    # databases
    "postgres", "postgresql", "mysql", "mongodb", "redis", "elasticsearch", "dynamodb",
    # mobile
    "android", "ios", "react native", "flutter",
    # soft / misc
    "communication", "stakeholder management", "agile", "scrum",
}

_SECTION_RX = re.compile(
    r"^\s*(skills?|technical\s+skills?|tech\s+stack|tools?|education|experience|"
    r"work\s+experience|professional\s+experience|projects?|summary|profile|"
    r"objective|certifications?)\s*[:\-]?\s*$",
    re.IGNORECASE,
)
_EMAIL_RX = re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}")
_PHONE_RX = re.compile(r"(?:\+?\d{1,3}[\s\-.])?\(?\d{2,4}\)?[\s\-.]?\d{3,4}[\s\-.]?\d{3,4}")
_YEAR_RX = re.compile(r"\b(19[8-9]\d|20\d{2})\b")
_YEARS_EXP_RX = re.compile(
    r"\b(\d{1,2}(?:\.\d)?)\s*(?:\+?)\s*(?:years?|yrs?)\b", re.IGNORECASE
)


# ---------------------------------------------------------------------------- #
# File-format extraction
# ---------------------------------------------------------------------------- #
def extract_pdf_text(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("pypdf is not installed") from exc
    try:
        reader = PdfReader(io.BytesIO(data))
        pages = []
        for p in reader.pages:
            try:
                pages.append(p.extract_text() or "")
            except Exception:
                continue
        return "\n".join(pages)
    except Exception as exc:
        logger.warning("PDF extraction failed: %s", exc)
        return ""


def extract_docx_text(data: bytes) -> str:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("python-docx is not installed") from exc
    try:
        doc = docx.Document(io.BytesIO(data))
        lines = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        lines.append(cell.text)
        return "\n".join(lines)
    except Exception as exc:
        logger.warning("DOCX extraction failed: %s", exc)
        return ""


def extract_text(filename: str, content_type: str, data: bytes) -> str:
    name = (filename or "").lower()
    ct = (content_type or "").lower()
    if name.endswith(".pdf") or "pdf" in ct:
        return extract_pdf_text(data)
    if name.endswith(".docx") or "wordprocessingml" in ct:
        return extract_docx_text(data)
    if name.endswith(".txt") or ct.startswith("text/"):
        try:
            return data.decode("utf-8", errors="ignore")
        except Exception:
            return ""
    raise ValueError("Unsupported file type. Upload a PDF, DOCX, or TXT resume.")


# ---------------------------------------------------------------------------- #
# Heuristic section parsing
# ---------------------------------------------------------------------------- #
def _iter_sections(text: str) -> Iterable[Tuple[str, List[str]]]:
    section: Optional[str] = None
    bucket: List[str] = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        m = _SECTION_RX.match(line)
        if m:
            if section and bucket:
                yield section, bucket
            section = m.group(1).lower()
            bucket = []
            continue
        if section:
            bucket.append(line)
    if section and bucket:
        yield section, bucket


def _normalize_skill(s: str) -> str:
    return s.strip().lower().rstrip(".,;:")


def _mine_skills(text: str) -> List[str]:
    """Return a deduped, lower-cased list of skills found via lexicon + section heuristic."""
    found: List[str] = []
    lower = text.lower()
    for skill in SKILLS_LEXICON:
        # Word-boundary match where possible.
        pattern = r"(?:(?<=\W)|^)" + re.escape(skill) + r"(?=\W|$)"
        if re.search(pattern, lower):
            found.append(skill)

    # Pick up comma/bullet-separated tokens under a "skills" section.
    for section, lines in _iter_sections(text):
        if "skill" not in section and "tool" not in section and "tech" not in section:
            continue
        blob = " ".join(lines)
        parts = re.split(r"[,•·\u2022\|/]+", blob)
        for p in parts:
            token = _normalize_skill(p)
            if 2 <= len(token) <= 40 and token not in found:
                found.append(token)
    # Dedupe preserving order.
    seen = set()
    out = []
    for s in found:
        if s not in seen:
            seen.add(s)
            out.append(s)
    return out[:80]


def _mine_emails(text: str) -> List[str]:
    return sorted({m.group(0).lower() for m in _EMAIL_RX.finditer(text)})


def _mine_phones(text: str) -> List[str]:
    # Filter false positives (lines without enough digits).
    out = []
    for m in _PHONE_RX.finditer(text):
        s = m.group(0)
        digits = re.sub(r"\D", "", s)
        if 7 <= len(digits) <= 15:
            out.append(s.strip())
    return list(dict.fromkeys(out))[:5]


def _mine_section_lines(text: str, keywords: Iterable[str], limit: int = 20) -> List[str]:
    out: List[str] = []
    kws = tuple(keywords)
    for section, lines in _iter_sections(text):
        if any(k in section for k in kws):
            out.extend(lines)
    return out[:limit]


def _estimate_years_experience(text: str) -> Optional[float]:
    matches = [float(m.group(1)) for m in _YEARS_EXP_RX.finditer(text)]
    if matches:
        return max(matches)
    years = sorted({int(m.group(0)) for m in _YEAR_RX.finditer(text)})
    if len(years) >= 2:
        span = years[-1] - years[0]
        if 0 < span <= 50:
            return float(span)
    return None


def parse_resume_text(text: str) -> dict:
    """Return a dict that matches the `ParsedResume` shape."""
    text = text or ""
    return {
        "skills": _mine_skills(text),
        "emails": _mine_emails(text),
        "phones": _mine_phones(text),
        "education": _mine_section_lines(text, ("education",)),
        "experiences": _mine_section_lines(text, ("experience", "work")),
        "years_experience": _estimate_years_experience(text),
    }
